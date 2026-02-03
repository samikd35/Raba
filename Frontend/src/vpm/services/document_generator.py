"""
Document Generator Service for VPM
Generates professional PDF and Word documents with Yuba branding.
"""

from typing import List, Dict, Any, Optional
from datetime import datetime
from io import BytesIO
import os

# PDF Generation
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter, A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Table, TableStyle, Paragraph, Spacer, Image, PageBreak
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT, TA_JUSTIFY
from reportlab.pdfgen import canvas

# Word Generation
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH


class YubaDocumentGenerator:
    """
    Professional document generator with Yuba branding.
    Supports PDF and Word formats.
    """
    
    # Yuba Brand Colors
    BRAND_BLUE = colors.HexColor('#244694')
    BRAND_LIGHT_BLUE = colors.HexColor('#3B82F6')
    BRAND_GREEN = colors.HexColor('#10B981')
    BRAND_RED = colors.HexColor('#EF4444')
    BRAND_YELLOW = colors.HexColor('#F59E0B')
    
    # Component Type Colors
    JTBD_COLOR = colors.HexColor('#10B981')  # Green
    PAIN_COLOR = colors.HexColor('#EF4444')   # Red
    GAIN_COLOR = colors.HexColor('#3B82F6')   # Blue
    
    def __init__(self):
        """Initialize document generator"""
        self.logo_path = self._get_logo_path()
    
    def _get_logo_path(self) -> Optional[str]:
        """Get path to Yuba logo"""
        # Try multiple possible locations - prefer yubanow_logo.jpeg
        possible_paths = [
            os.path.join(os.path.dirname(__file__), '..', 'yubanow_logo.jpeg'),
            os.path.join(os.path.dirname(__file__), '..', 'assets', 'yubanow_logo.jpeg'),
            os.path.join(os.path.dirname(__file__), '..', 'assets', 'yuba_logo.png'),
            os.path.join(os.path.dirname(__file__), '..', '..', 'assets', 'yuba_logo.png'),
            '/app/assets/yubanow_logo.jpeg',  # Docker path
            '/app/assets/yuba_logo.png',  # Docker path fallback
        ]
        
        for path in possible_paths:
            if os.path.exists(path):
                return path
        
        return None
    
    def _get_component_color(self, component_type: str) -> colors.Color:
        """Get color for component type"""
        component_type_lower = component_type.lower()
        if component_type_lower == 'jtbd':
            return self.JTBD_COLOR
        elif component_type_lower == 'pain':
            return self.PAIN_COLOR
        elif component_type_lower == 'gain':
            return self.GAIN_COLOR
        else:
            return colors.grey
    
    def _get_component_label(self, component_type: str) -> str:
        """Get display label for component type"""
        component_type_lower = component_type.lower()
        if component_type_lower == 'jtbd':
            return 'JTBD'
        elif component_type_lower == 'pain':
            return 'Pain'
        elif component_type_lower == 'gain':
            return 'Gain'
        else:
            return component_type
    
    async def generate_questionnaires_pdf(
        self,
        questionnaires: List[Dict[str, Any]],
        project_name: str,
        project_id: str
    ) -> BytesIO:
        """
        Generate PDF document with questionnaires.
        
        Args:
            questionnaires: List of questionnaire items
            project_name: Name of the project
            project_id: Project ID
            
        Returns:
            BytesIO: PDF file content
        """
        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Container for document elements
        elements = []
        
        # Styles
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=self.BRAND_BLUE,
            spaceAfter=30,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        subtitle_style = ParagraphStyle(
            'CustomSubtitle',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.grey,
            spaceAfter=20,
            alignment=TA_CENTER
        )
        
        persona_style = ParagraphStyle(
            'PersonaStyle',
            parent=styles['Heading2'],
            fontSize=16,
            textColor=self.BRAND_BLUE,
            spaceAfter=12,
            spaceBefore=20,
            fontName='Helvetica-Bold'
        )
        
        question_number_style = ParagraphStyle(
            'QuestionNumber',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.grey,
            spaceAfter=6
        )
        
        question_text_style = ParagraphStyle(
            'QuestionText',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.black,
            spaceAfter=8,
            leading=16
        )
        
        metadata_style = ParagraphStyle(
            'Metadata',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.grey,
            spaceAfter=20
        )
        
        # Add logo if available
        if self.logo_path and os.path.exists(self.logo_path):
            try:
                logo = Image(self.logo_path, width=1.5*inch, height=0.5*inch)
                elements.append(logo)
                elements.append(Spacer(1, 12))
            except Exception as e:
                print(f"Warning: Could not load logo: {e}")
        
        # Title
        elements.append(Paragraph("Interview Questions", title_style))
        
        # Project info
        elements.append(Paragraph(f"Project: {project_name}", subtitle_style))
        elements.append(Spacer(1, 20))
        
        # Group questionnaires by persona
        personas_dict = {}
        for q in questionnaires:
            persona_name = q.get('persona_name', 'Unknown Persona')
            if persona_name not in personas_dict:
                personas_dict[persona_name] = []
            personas_dict[persona_name].append(q)
        
        # Add questionnaires by persona
        for persona_idx, (persona_name, persona_questions) in enumerate(personas_dict.items()):
            # Persona header
            elements.append(Paragraph(f"{persona_name}", persona_style))
            elements.append(Paragraph(
                f"{len(persona_questions)} interview questions",
                metadata_style
            ))
            
            # Add questions
            for q_idx, question in enumerate(persona_questions, 1):
                # Question number and component type badge
                component_type = question.get('component_type', 'unknown')
                component_label = self._get_component_label(component_type)
                component_color = self._get_component_color(component_type)
                
                # Create styled question header paragraph
                question_header_style = ParagraphStyle(
                    'QuestionHeader',
                    parent=styles['Normal'],
                    fontSize=11,
                    textColor=colors.black,
                    spaceAfter=6,
                    spaceBefore=6
                )
                
                # Create the header with question number and colored badge
                header_text = f"{q_idx}. <font color='{component_color.hexval()}' size='10'><b>Validates: {component_label}</b></font>"
                elements.append(Paragraph(header_text, question_header_style))
                
                # Question text
                question_text = question.get('text', 'No question text')
                elements.append(Paragraph(question_text, question_text_style))
                
                # Separator line
                elements.append(Spacer(1, 12))
                
                # Add line separator between questions
                if q_idx < len(persona_questions):
                    line_data = [['']]
                    line_table = Table(line_data, colWidths=[6.5*inch])
                    line_table.setStyle(TableStyle([
                        ('LINEABOVE', (0, 0), (-1, 0), 1, colors.lightgrey),
                    ]))
                    elements.append(line_table)
                    elements.append(Spacer(1, 12))
            
            # Page break between personas (except last one)
            if persona_idx < len(personas_dict) - 1:
                elements.append(PageBreak())
        
        # Footer info
        elements.append(Spacer(1, 30))
        footer_style = ParagraphStyle(
            'Footer',
            parent=styles['Normal'],
            fontSize=8,
            textColor=colors.grey,
            alignment=TA_CENTER
        )
        elements.append(Paragraph(
            f"© {datetime.now().year} Yuba | Total Questions: {len(questionnaires)}",
            footer_style
        ))
        
        # Build PDF
        doc.build(elements)
        buffer.seek(0)
        return buffer
    
    async def generate_questionnaires_docx(
        self,
        questionnaires: List[Dict[str, Any]],
        project_name: str,
        project_id: str
    ) -> BytesIO:
        """
        Generate Word document with questionnaires.
        
        Args:
            questionnaires: List of questionnaire items
            project_name: Name of the project
            project_id: Project ID
            
        Returns:
            BytesIO: Word file content
        """
        doc = Document()
        
        # Set default font
        style = doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(11)
        
        # Add logo if available
        if self.logo_path and os.path.exists(self.logo_path):
            try:
                doc.add_picture(self.logo_path, width=Inches(2.0))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                print(f"Warning: Could not load logo: {e}")
        
        # Title
        title = doc.add_heading('Interview Questions', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_run = title.runs[0]
        title_run.font.color.rgb = RGBColor(36, 70, 148)  # Yuba blue
        
        # Project info
        project_info = doc.add_paragraph()
        project_info.alignment = WD_ALIGN_PARAGRAPH.CENTER
        project_run = project_info.add_run(f"Project: {project_name}")
        project_run.font.size = Pt(12)
        project_run.font.color.rgb = RGBColor(128, 128, 128)
        
        doc.add_paragraph()  # Spacer
        
        # Group questionnaires by persona
        personas_dict = {}
        for q in questionnaires:
            persona_name = q.get('persona_name', 'Unknown Persona')
            if persona_name not in personas_dict:
                personas_dict[persona_name] = []
            personas_dict[persona_name].append(q)
        
        # Add questionnaires by persona
        for persona_idx, (persona_name, persona_questions) in enumerate(personas_dict.items()):
            # Persona header
            persona_heading = doc.add_heading(persona_name, 1)
            persona_run = persona_heading.runs[0]
            persona_run.font.color.rgb = RGBColor(36, 70, 148)
            
            # Persona subtitle
            subtitle = doc.add_paragraph(f"{len(persona_questions)} interview questions")
            subtitle_run = subtitle.runs[0]
            subtitle_run.font.size = Pt(10)
            subtitle_run.font.color.rgb = RGBColor(128, 128, 128)
            
            # Add questions
            for q_idx, question in enumerate(persona_questions, 1):
                # Question header
                question_para = doc.add_paragraph()
                
                # Question number
                q_num_run = question_para.add_run(f"{q_idx}.  ")
                q_num_run.font.size = Pt(10)
                q_num_run.font.color.rgb = RGBColor(128, 128, 128)
                
                # Component type badge
                component_type = question.get('component_type', 'unknown')
                component_label = self._get_component_label(component_type)
                badge_run = question_para.add_run(f"Validates: {component_label}")
                badge_run.font.size = Pt(10)
                badge_run.font.bold = True
                
                # Set color based on component type
                if component_type.lower() == 'jtbd':
                    badge_run.font.color.rgb = RGBColor(16, 185, 129)  # Green
                elif component_type.lower() == 'pain':
                    badge_run.font.color.rgb = RGBColor(239, 68, 68)   # Red
                elif component_type.lower() == 'gain':
                    badge_run.font.color.rgb = RGBColor(59, 130, 246)  # Blue
                
                # Question text
                question_text = question.get('text', 'No question text')
                text_para = doc.add_paragraph(question_text)
                text_para.paragraph_format.space_before = Pt(6)
                text_para.paragraph_format.space_after = Pt(6)
                
                # Add spacing between questions
                if q_idx < len(persona_questions):
                    doc.add_paragraph()
            
            # Page break between personas (except last one)
            if persona_idx < len(personas_dict) - 1:
                doc.add_page_break()
        
        # Footer
        doc.add_paragraph()
        footer = doc.add_paragraph()
        footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer.add_run(
            f"© {datetime.now().year} Yuba | Total Questions: {len(questionnaires)}"
        )
        footer_run.font.size = Pt(8)
        footer_run.font.color.rgb = RGBColor(128, 128, 128)
        
        # Save to buffer
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer

    # ============================================================
    # PROBLEM VALIDATION REPORT PDF GENERATION
    # ============================================================

    def _create_header_footer(self, canvas_obj: canvas.Canvas, doc: SimpleDocTemplate):
        """Add branded header and footer to each page."""
        page_width, page_height = letter
        
        # ===== HEADER =====
        # Header background bar
        canvas_obj.setFillColor(self.BRAND_BLUE)
        canvas_obj.rect(0, page_height - 50, page_width, 50, fill=True, stroke=False)
        
        # Logo in header (white text fallback if logo not available)
        if self.logo_path and os.path.exists(self.logo_path):
            try:
                canvas_obj.drawImage(
                    self.logo_path, 
                    40, page_height - 40,
                    width=80, height=25,
                    preserveAspectRatio=True,
                    mask='auto'
                )
            except Exception:
                # Fallback to text
                canvas_obj.setFillColor(colors.white)
                canvas_obj.setFont('Helvetica-Bold', 16)
                canvas_obj.drawString(40, page_height - 35, "YUBA")
        else:
            canvas_obj.setFillColor(colors.white)
            canvas_obj.setFont('Helvetica-Bold', 16)
            canvas_obj.drawString(40, page_height - 35, "YUBA")
        
        # "Problem Validation Report" text in header
        canvas_obj.setFillColor(colors.white)
        canvas_obj.setFont('Helvetica', 10)
        canvas_obj.drawRightString(page_width - 40, page_height - 35, "Problem Validation Report")
        
        # ===== FOOTER =====
        # Footer background bar
        canvas_obj.setFillColor(self.BRAND_BLUE)
        canvas_obj.rect(0, 0, page_width, 35, fill=True, stroke=False)
        
        # Footer text - yubanow.com link
        canvas_obj.setFillColor(colors.white)
        canvas_obj.setFont('Helvetica', 9)
        canvas_obj.drawString(40, 14, f"© {datetime.now().year} Yuba | yubanow.com")
        
        # Page number
        page_num = canvas_obj.getPageNumber()
        canvas_obj.drawRightString(page_width - 40, 14, f"Page {page_num}")

    def _on_first_page(self, canvas_obj: canvas.Canvas, doc: SimpleDocTemplate):
        """First page header/footer (same as later pages)."""
        self._create_header_footer(canvas_obj, doc)

    def _on_later_pages(self, canvas_obj: canvas.Canvas, doc: SimpleDocTemplate):
        """Later pages header/footer."""
        self._create_header_footer(canvas_obj, doc)

    async def generate_pv_report_pdf(
        self,
        report_data: Dict[str, Any],
        report_title: str
    ) -> BytesIO:
        """
        Generate a professional PDF for Problem Validation Report with Yuba branding.
        
        Args:
            report_data: Dictionary containing PV report sections:
                - title: Report title
                - executive_summary: Executive summary text
                - industry_analysis: Industry analysis with markdown subsections
                - challenges_analysis: PESTEL analysis with markdown subsections
                - recommendations: Recommendations text
                - sources: List of references [{number, source_url, source_title}]
                - industry: Industry field
                - geography: Geography field
                - created_at: Creation timestamp
            report_title: Override title for the report
            
        Returns:
            BytesIO: PDF file content buffer
        """
        buffer = BytesIO()
        
        # Store sources for clickable references
        self._sources = report_data.get('sources', [])
        self._source_url_map = {s.get('number', i+1): s.get('source_url', '') for i, s in enumerate(self._sources)}
        
        # Create document with custom margins for header/footer
        doc = SimpleDocTemplate(
            buffer,
            pagesize=letter,
            rightMargin=50,
            leftMargin=50,
            topMargin=70,  # Space for header
            bottomMargin=55  # Space for footer
        )
        
        # Container for document elements
        elements = []
        
        # Get styles
        styles = getSampleStyleSheet()
        
        # ===== CUSTOM STYLES =====
        
        # Main Title Style - "Problem Validation Report"
        title_style = ParagraphStyle(
            'PVReportTitle',
            parent=styles['Heading1'],
            fontSize=22,
            textColor=self.BRAND_BLUE,
            spaceAfter=8,
            spaceBefore=20,
            alignment=TA_CENTER,
            fontName='Helvetica-Bold'
        )
        
        # Report metadata style
        metadata_style = ParagraphStyle(
            'PVMetadata',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.grey,
            spaceAfter=4,
            alignment=TA_CENTER
        )
        
        # Body text style
        body_style = ParagraphStyle(
            'PVBodyText',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.black,
            spaceAfter=8,
            leading=14,
            alignment=TA_JUSTIFY
        )
        
        # Executive summary style - clean, readable, no ugly box
        exec_summary_style = ParagraphStyle(
            'PVExecSummary',
            parent=styles['Normal'],
            fontSize=10,
            textColor=colors.HexColor('#374151'),  # Dark gray for readability
            spaceAfter=10,
            leading=15,
            alignment=TA_JUSTIFY,
            firstLineIndent=0,
        )
        
        # Subsection heading style (## headers)
        subsection_style = ParagraphStyle(
            'PVSubsection',
            parent=styles['Heading3'],
            fontSize=12,
            textColor=self.BRAND_BLUE,
            spaceBefore=14,
            spaceAfter=6,
            fontName='Helvetica-Bold',
        )
        
        # Sub-subsection heading style (### headers)
        subsubsection_style = ParagraphStyle(
            'PVSubSubsection',
            parent=styles['Heading4'],
            fontSize=10,
            textColor=colors.HexColor('#1E40AF'),  # Slightly lighter blue
            spaceBefore=10,
            spaceAfter=4,
            fontName='Helvetica-Bold',
        )
        
        # Source/reference style
        source_style = ParagraphStyle(
            'PVSource',
            parent=styles['Normal'],
            fontSize=9,
            textColor=colors.HexColor('#4B5563'),
            spaceAfter=4,
            leading=12,
            leftIndent=20,
            firstLineIndent=-20,
        )
        
        # ===== DOCUMENT CONTENT =====
        
        # Title - Always "Problem Validation Report"
        elements.append(Paragraph("Problem Validation Report", title_style))
        
        # Metadata line
        created_at = report_data.get('created_at', '')
        if created_at:
            try:
                if isinstance(created_at, str):
                    from datetime import datetime as dt
                    created_date = dt.fromisoformat(created_at.replace('Z', '+00:00'))
                    created_str = created_date.strftime('%B %d, %Y')
                else:
                    created_str = created_at.strftime('%B %d, %Y')
            except Exception:
                created_str = str(created_at)[:10]
        else:
            created_str = datetime.now().strftime('%B %d, %Y')
        
        elements.append(Paragraph(f"Generated: {created_str}", metadata_style))
        elements.append(Spacer(1, 15))
        
        # Horizontal line separator
        line_data = [['']]
        line_table = Table(line_data, colWidths=[6.5*inch])
        line_table.setStyle(TableStyle([
            ('LINEABOVE', (0, 0), (-1, 0), 2, self.BRAND_BLUE),
        ]))
        elements.append(line_table)
        elements.append(Spacer(1, 15))
        
        # ===== 1. EXECUTIVE SUMMARY =====
        exec_summary = report_data.get('executive_summary', '')
        if exec_summary and exec_summary.strip():
            elements.append(self._create_section_header("1. Executive Summary", styles))
            # Format with clickable citations and clean text
            formatted_summary = self._format_content_with_links(exec_summary)
            elements.append(Paragraph(formatted_summary, exec_summary_style))
            elements.append(Spacer(1, 10))
        
        # ===== 2. INDUSTRY ANALYSIS (with dynamic subsections) =====
        industry_analysis = report_data.get('industry_analysis', '')
        if industry_analysis and industry_analysis.strip():
            elements.append(PageBreak())  # Start on new page
            elements.append(self._create_section_header("2. Industry Analysis", styles))
            # Parse and add subsections dynamically
            self._add_markdown_sections(elements, industry_analysis, body_style, subsection_style, subsubsection_style)
            elements.append(Spacer(1, 10))
        
        # ===== 3. PESTEL ANALYSIS (with dynamic subsections) =====
        challenges_analysis = report_data.get('challenges_analysis', '')
        if challenges_analysis and challenges_analysis.strip():
            elements.append(PageBreak())  # Start on new page
            elements.append(self._create_section_header("3. PESTEL Analysis & Market Challenges", styles))
            # Parse and add subsections dynamically
            self._add_markdown_sections(elements, challenges_analysis, body_style, subsection_style, subsubsection_style)
            elements.append(Spacer(1, 10))
        
        # ===== 4. STRATEGIC RECOMMENDATIONS =====
        recommendations = report_data.get('recommendations', '')
        if recommendations and recommendations.strip():
            elements.append(PageBreak())  # Start on new page
            elements.append(self._create_section_header("4. Strategic Recommendations", styles))
            formatted_recs = self._format_content_with_links(recommendations)
            elements.append(Paragraph(formatted_recs, body_style))
            elements.append(Spacer(1, 10))
        
        # ===== 5. SOURCES (with clickable links) =====
        sources = report_data.get('sources', [])
        if sources and len(sources) > 0:
            elements.append(PageBreak())  # Start on new page
            elements.append(self._create_section_header("5. Sources & References", styles))
            elements.append(Spacer(1, 5))
            
            for source in sources:
                source_num = source.get('number', '')
                source_title = source.get('source_title', '') or source.get('title', 'Untitled')
                source_url = source.get('source_url', '') or source.get('url', '')
                
                if source_url:
                    # Clickable link
                    source_text = f"[{source_num}] <b>{source_title}</b><br/><link href=\"{source_url}\" color=\"blue\">{source_url[:80]}{'...' if len(source_url) > 80 else ''}</link>"
                else:
                    source_text = f"[{source_num}] <b>{source_title}</b>"
                
                elements.append(Paragraph(source_text, source_style))
        
        # ===== FINAL SPACER =====
        elements.append(Spacer(1, 30))
        
        # Build PDF with custom header/footer
        doc.build(
            elements,
            onFirstPage=self._on_first_page,
            onLaterPages=self._on_later_pages
        )
        
        buffer.seek(0)
        return buffer
    
    def _add_markdown_sections(self, elements, content: str, body_style, subsection_style, subsubsection_style):
        """Parse markdown-like content and add sections dynamically."""
        import re
        
        lines = content.split('\n')
        current_text = []
        
        for line in lines:
            # Check for ## header (subsection)
            if line.strip().startswith('## '):
                # Flush current text
                if current_text:
                    text = '\n'.join(current_text)
                    formatted = self._format_content_with_links(text)
                    if formatted.strip():
                        elements.append(Paragraph(formatted, body_style))
                    current_text = []
                
                # Add subsection header
                header_text = line.strip()[3:].strip()
                elements.append(Paragraph(header_text, subsection_style))
            
            # Check for ### header (sub-subsection)
            elif line.strip().startswith('### '):
                # Flush current text
                if current_text:
                    text = '\n'.join(current_text)
                    formatted = self._format_content_with_links(text)
                    if formatted.strip():
                        elements.append(Paragraph(formatted, body_style))
                    current_text = []
                
                # Add sub-subsection header
                header_text = line.strip()[4:].strip()
                elements.append(Paragraph(header_text, subsubsection_style))
            
            else:
                # Regular content line
                current_text.append(line)
        
        # Flush remaining text
        if current_text:
            text = '\n'.join(current_text)
            formatted = self._format_content_with_links(text)
            if formatted.strip():
                elements.append(Paragraph(formatted, body_style))
    
    def _format_content_with_links(self, content: str) -> str:
        """Format content with clickable citation links and markdown formatting."""
        if not content:
            return ""
        
        import re
        
        # Replace markdown bold **text** with PDF bold tags
        content = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', content)
        
        # Replace markdown italic *text* with PDF italic tags
        content = re.sub(r'(?<!\*)\*([^*]+)\*(?!\*)', r'<i>\1</i>', content)
        content = re.sub(r'_([^_]+)_', r'<i>\1</i>', content)
        
        # Replace markdown bullet points
        content = re.sub(r'^\s*[-*]\s+', '• ', content, flags=re.MULTILINE)
        
        # Make citation numbers clickable [X] -> link to source
        def make_citation_link(match):
            citation_num = match.group(1)
            try:
                num = int(citation_num)
                url = self._source_url_map.get(num, '')
                if url:
                    return f'<link href="{url}" color="blue">[{citation_num}]</link>'
            except (ValueError, AttributeError):
                pass
            return match.group(0)
        
        content = re.sub(r'\[(\d+)\]', make_citation_link, content)
        
        # Replace newlines with proper breaks
        content = content.replace('\n\n', '<br/><br/>')
        content = content.replace('\n', '<br/>')
        
        return content

    def _create_section_header(self, title: str, styles) -> Table:
        """Create a styled section header with colored left border."""
        # Create a table with colored left border effect
        header_style = ParagraphStyle(
            'SectionHeader',
            parent=styles['Heading2'],
            fontSize=13,
            textColor=self.BRAND_BLUE,
            fontName='Helvetica-Bold',
            leftIndent=10,
        )
        
        header_para = Paragraph(title, header_style)
        
        # Create table with left border
        data = [[header_para]]
        table = Table(data, colWidths=[6.3*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, -1), colors.HexColor('#F0F4F8')),
            ('LEFTPADDING', (0, 0), (-1, -1), 12),
            ('RIGHTPADDING', (0, 0), (-1, -1), 12),
            ('TOPPADDING', (0, 0), (-1, -1), 8),
            ('BOTTOMPADDING', (0, 0), (-1, -1), 8),
            ('LINEBELOW', (0, 0), (-1, -1), 0, colors.white),
            # Left colored border effect
            ('LINEBEFORE', (0, 0), (0, -1), 4, self.BRAND_BLUE),
        ]))
        
        return table

    def _format_content_for_pdf(self, content: str) -> str:
        """Format markdown-like content for PDF paragraphs."""
        if not content:
            return ""
        
        # Replace markdown bold **text** with PDF bold tags
        import re
        content = re.sub(r'\*\*([^*]+)\*\*', r'<b>\1</b>', content)
        
        # Replace markdown italic *text* or _text_ with PDF italic tags
        content = re.sub(r'(?<!\*)\*([^*]+)\*(?!\*)', r'<i>\1</i>', content)
        content = re.sub(r'_([^_]+)_', r'<i>\1</i>', content)
        
        # Replace markdown headers with bold
        content = re.sub(r'^#+\s*(.+)$', r'<b>\1</b>', content, flags=re.MULTILINE)
        
        # Replace markdown bullet points with proper formatting
        content = re.sub(r'^\s*[-*]\s+', '• ', content, flags=re.MULTILINE)
        
        # Replace newlines with proper breaks
        content = content.replace('\n\n', '<br/><br/>')
        content = content.replace('\n', '<br/>')
        
        return content

    def _dict_to_text(self, data: Dict[str, Any], indent: int = 0) -> str:
        """Convert dictionary data to formatted text."""
        if not data:
            return "No data available"
        
        lines = []
        indent_str = "&nbsp;" * (indent * 4)
        
        for key, value in data.items():
            # Format key nicely
            formatted_key = key.replace('_', ' ').title()
            
            if isinstance(value, dict):
                lines.append(f"{indent_str}<b>{formatted_key}:</b>")
                lines.append(self._dict_to_text(value, indent + 1))
            elif isinstance(value, list):
                lines.append(f"{indent_str}<b>{formatted_key}:</b>")
                for item in value:
                    if isinstance(item, dict):
                        lines.append(self._dict_to_text(item, indent + 1))
                    else:
                        lines.append(f"{indent_str}&nbsp;&nbsp;• {item}")
            else:
                lines.append(f"{indent_str}<b>{formatted_key}:</b> {value}")
        
        return "<br/>".join(lines)
